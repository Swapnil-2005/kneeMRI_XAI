import streamlit as st
import cv2
import numpy as np
import tensorflow as tf
from PIL import Image
import matplotlib.cm as cm
from ultralytics import YOLO
from docx import Document
from docx.shared import Inches
from datetime import datetime
import tempfile

# ---------------- CONFIG ----------------
st.set_page_config(page_title="NeuroRad Vision", layout="wide")

YOLO_MODEL = "acl.pt"
CNN_WEIGHTS = "model.weights.h5"
IMG_SIZE = (224, 224)

# ---------------- LOAD MODELS ----------------
@st.cache_resource
def load_models():
    yolo = YOLO(YOLO_MODEL)

    base = tf.keras.applications.EfficientNetB0(
        include_top=False,
        weights=None,
        input_shape=(224,224,3)
    )
    x = tf.keras.layers.GlobalAveragePooling2D()(base.output)
    x = tf.keras.layers.Dense(128, activation="relu")(x)
    x = tf.keras.layers.Dropout(0.6)(x)
    out = tf.keras.layers.Dense(2, activation="softmax")(x)

    cnn = tf.keras.Model(base.input, out)
    cnn.load_weights(CNN_WEIGHTS)

    return yolo, cnn

yolo_model, cnn_model = load_models()

# ---------------- FUNCTIONS ----------------

def preprocess(img):
    img = cv2.resize(img, IMG_SIZE) / 255.0
    return np.expand_dims(img, 0)

def crop_yolo(res, img):
    boxes = res.boxes.xyxy.cpu().numpy()
    if len(boxes) == 0:
        return None
    x1, y1, x2, y2 = map(int, boxes[0])
    return img[y1:y2, x1:x2]

# ---------------- CLINICAL GRADCAM ----------------

def make_gradcam_heatmap(img_array, model, last_conv_layer_name):
    grad_model = tf.keras.models.Model(
        model.input,
        [model.get_layer(last_conv_layer_name).output, model.output]
    )

    with tf.GradientTape() as tape:
        conv_output, preds = grad_model(img_array)
        class_idx = tf.argmax(preds[0])
        loss = preds[:, class_idx]

    grads = tape.gradient(loss, conv_output)
    pooled_grads = tf.reduce_mean(grads, axis=(0,1,2))

    conv_output = conv_output[0]
    heatmap = conv_output @ pooled_grads[..., tf.newaxis]
    heatmap = tf.squeeze(heatmap)

    # ---------------- CLINICAL FILTERING ----------------
    heatmap = tf.maximum(heatmap, 0)

    # normalize safely
    heatmap /= (tf.reduce_max(heatmap) + 1e-8)

    heatmap = heatmap.numpy()

    # 🔥 Smooth (remove noise)
    heatmap = cv2.GaussianBlur(heatmap, (7,7), 0)

    # 🔥 Threshold (focus only important region)
    heatmap[heatmap < 0.4] = 0

    return heatmap

def overlay_gradcam(original_img, heatmap, alpha=0.5):
    # Convert heatmap to 0-255
    heatmap = np.uint8(255 * heatmap)

    # Apply colormap
    jet = cm.get_cmap("jet")
    jet_colors = jet(np.arange(256))[:, :3]
    jet_heatmap = jet_colors[heatmap]

    # Convert to numpy (NOT PIL)
    jet_heatmap = np.uint8(jet_heatmap * 255)

    # Resize both to SAME SIZE
    jet_heatmap = cv2.resize(jet_heatmap, (224, 224))
    base = cv2.resize(original_img, (224, 224))

    # Ensure same dtype
    jet_heatmap = jet_heatmap.astype(np.float32)
    base = base.astype(np.float32)

    # Overlay
    overlay = jet_heatmap * alpha + base

    # Normalize back
    overlay = np.clip(overlay, 0, 255).astype(np.uint8)

    return overlay

# ---------------- REPORT ----------------

def generate_report(original, detection, gradcam):
    doc = Document()

    doc.add_heading("NeuroRad Vision - Clinical Report", 0)

    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("Findings", level=1)
    doc.add_paragraph("Automated analysis performed on knee MRI.")

    doc.add_heading("Images", level=1)

    tmp1 = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    tmp2 = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    tmp3 = tempfile.NamedTemporaryFile(delete=False, suffix=".png")

    Image.fromarray(original).save(tmp1.name)
    Image.fromarray(detection).save(tmp2.name)
    Image.fromarray(gradcam).save(tmp3.name)

    doc.add_paragraph("Original MRI:")
    doc.add_picture(tmp1.name, width=Inches(3))

    doc.add_paragraph("Detection:")
    doc.add_picture(tmp2.name, width=Inches(3))

    doc.add_paragraph("Grad-CAM (Explainability):")
    doc.add_picture(tmp3.name, width=Inches(3))

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "Regions of interest detected using deep learning. "
        "Highlighted zones indicate model attention for diagnosis support."
    )

    report_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(report_path)

    return report_path

# ---------------- UI ----------------

st.title("🧠 NeuroRad Vision")

uploaded = st.file_uploader("Upload Knee MRI", type=["jpg","png","jpeg"])

if uploaded:
    img = np.array(Image.open(uploaded).convert("RGB"))
    img_cv = cv2.cvtColor(img, cv2.COLOR_RGB2BGR)

    # YOLO
    results = yolo_model(img_cv)
    yolo_img = results[0].plot(conf=False)
    yolo_img = cv2.cvtColor(yolo_img, cv2.COLOR_BGR2RGB)

    crop = crop_yolo(results[0], img_cv)

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Original MRI")
        st.image(img)

    with col2:
        st.subheader("Detection")
        st.image(yolo_img)

    if crop is not None:
        crop_rgb = cv2.cvtColor(crop, cv2.COLOR_BGR2RGB)

        img_array = preprocess(crop_rgb)

        heatmap = make_gradcam_heatmap(
            img_array,
            cnn_model,
            last_conv_layer_name="top_conv"
        )

        gradcam_img = overlay_gradcam(crop_rgb, heatmap)

        st.subheader("🔍 Grad-CAM (Clinical Focus)")
        st.image(gradcam_img, width=250)

        # ---------------- REPORT ----------------
        report_path = generate_report(img, yolo_img, gradcam_img)

        with open(report_path, "rb") as f:
            st.download_button(
                "📄 Download Clinical Report",
                f,
                file_name="NeuroRad_Report.docx"
            )

    else:
        st.warning("No ACL region detected")